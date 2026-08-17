from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

# 标题
title = doc.add_heading('课程设计报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('项目名称：Project 2: 基于价值迭代的倒立摆全局最优控制')
doc.add_paragraph('完成日期：2026年5月27日')
doc.add_paragraph('_' * 50)

# 读取 Markdown 文件
with open(r'd:\python\力学\课程设计报告_Project2_价值迭代倒立摆.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')  # 只去掉换行，保留行首空格用于缩进判断
    
    # ========== 空行处理 ==========
    if not line.strip():
        i += 1
        continue
    
    stripped = line.strip()
    
    # ========== 标题处理 ==========
    if stripped.startswith('# '):
        doc.add_heading(stripped[2:], level=1)
    elif stripped.startswith('## '):
        doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith('### '):
        doc.add_heading(stripped[4:], level=3)
    
    # ========== 表格处理 ==========
    elif stripped.startswith('|'):
        # 收集表格所有行
        table_lines = []
        while i < len(lines):
            current = lines[i].strip()
            if current.startswith('|'):
                table_lines.append(current)
                i += 1
            else:
                break
        
        # 过滤掉表头分隔行（如 |---|---|）
        data_lines = [l for l in table_lines if '---' not in l.replace('-', '') or not l.replace('|', '').replace('-', '').replace(':', '').strip() == '']
        # 更安全的过滤：只要包含 "---" 且没有其他内容的就是分隔行
        data_lines = []
        for l in table_lines:
            content = l.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')
            if content == '':
                continue  # 跳过分隔行
            data_lines.append(l)
        
        if data_lines:
            # 解析第一行确定列数
            first_row_cells = [c.strip() for c in data_lines[0].split('|')[1:-1]]
            num_cols = len(first_row_cells)
            
            # 创建表格（行数 = 数据行数）
            table = doc.add_table(rows=0, cols=num_cols)
            table.style = 'Table Grid'
            
            for row_data in data_lines:
                cells = [c.strip() for c in row_data.split('|')[1:-1]]
                row = table.add_row()
                for j, cell_text in enumerate(cells):
                    if j < num_cols:
                        row.cells[j].text = cell_text
        continue  # 已经 i += 1 过了，跳过最后的 i += 1
    
    # ========== 公式处理（行内 $...$） ==========
    elif '$' in stripped:
        p = doc.add_paragraph()
        # 用正则分割，保留分隔符
        parts = re.split(r'(\$[^$]+\$)', stripped)
        for part in parts:
            if not part:
                continue
            run = p.add_run()
            if part.startswith('$') and part.endswith('$'):
                formula = part[1:-1]
                run.text = formula
                # 可选：设置等宽字体模拟公式
                run.font.name = 'Cambria Math'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria Math')
            else:
                run.text = part
    
    # ========== 普通段落 ==========
    elif not stripped.startswith('---'):
        doc.add_paragraph(stripped)
    
    i += 1

doc.save('Project2_Report.docx')
print("✅ 已生成 Project2_Report.docx")